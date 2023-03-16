import streamlit as st
import re
import json
import base64
import time
import docx
from docx.shared import Inches
from docx2pdf import convert
import langchain
from langchain import PromptTemplate, LLMChain
from langchain.chat_models import ChatOpenAI
import pandas as pd

openai_api_key = st.secrets["OPENAI_API_KEY"]
pinecone_api_key = st.secrets["PINECONE_API_KEY"]
index_name = st.secrets["PINECONE_INDEX_NAME"]

# Main app function
def main():
    st.title("SynthScope: Automatic Systematic Review Tool")
    st.write(
        """
        SynthScope is an automatic systematic review tool that uses advanced AI LLMs to quickly and accurately summarize and synthesize large volumes of scientific literature that is imported by the user.
        With SynthScope, you can save time and reduce errors by automating the title/abstract, and full text screening stages of the literature review, enabling you to focus on the analysis and insights that matter most to your research.
        """
    )

    st.sidebar.title("SynthScope Features")
    feature = st.sidebar.selectbox(
        "Choose a feature",
        (
            "Select...",
            "Search Strategy Generation",
            "Database Search",
            "Title and Abstract Screening",
            "Full Text PDF Retrieval",
            "Full Text Screening",
            "Risk of Bias Screening",
            "Summary Generation",
            "Report Generation",
        ),
    )

    if feature == "Search Strategy Generation":
        search_strategy_generation()
    elif feature == "Database Search":
        database_search()
    elif feature == "Title and Abstract Screening":
        title_abstract_screening()
    elif feature == "Full Text PDF Retrieval":
        full_text_pdf_retrieval()
    elif feature == "Full Text Screening":
        full_text_screening()
    elif feature == "Risk of Bias Screening":
        risk_of_bias_screening()
    elif feature == "Summary Generation":
        summary_generation()
    elif feature == "Report Generation":
        report_generation()

def download_link(object_to_download, download_filename, download_link_text):
    """
    Generates a link to download the given object_to_download.
    """
    if isinstance(object_to_download, pd.DataFrame):
        object_to_download = object_to_download.to_csv(index=False)

    # Some strings <-> bytes conversions necessary here
    b64 = base64.b64encode(object_to_download.encode()).decode()

    return f'<a href="data:file/txt;base64,{b64}" download="{download_filename}">{download_link_text}</a>'

def search_strategy_generation():
    st.markdown("## Generate Search Strategy")
    user_input = st.text_input("Enter your research question:")
    if st.button("Generate"):
        if user_input:
            summary_dict = searchstrat_developer(user_input)

            docx_filename = str(summary_dict["File Title"]).strip('"').strip('.pdf') + ".docx"
            pdf_filename = str(summary_dict["File Title"]).strip('"') + ".pdf"

            # Display download links
            with open('summary_file.txt', 'r', encoding='utf-8') as file:
                summary_text = file.read()
            st.markdown(download_link(summary_text, "summary.txt", "Download Summary as TXT"), unsafe_allow_html=True)
            st.markdown(download_link(summary_dict, docx_filename, "Download Summary as DOCX"), unsafe_allow_html=True)
            st.markdown(download_link(summary_dict, pdf_filename, "Download Summary as PDF"), unsafe_allow_html=True)

        else:
            st.error("Please enter a research question.")


def search_strategy_generation():
    st.header("Search Strategy Generation")
    st.write("SynthScope will assist the user in generating a search strategy via the gpt-3.5-turbo OpenAI model and langchain.")
    st.markdown("## Generate Search Strategy")
    user_input = st.text_input("Enter your research question:")
    search_query = st.text_input("Enter your search query:")
    if st.button("Generate Search Strategy"):
        if user_input:
            output_placeholder = st.empty()
            summary = []
            summary_dict = {}

            # Ask the LLM for a new research question to-do: create a list of 5 different questions
            # Then ask the user which one they prefer or to keep the original question
            output_placeholder.write("Developing new research questions...")
            with open('perssonadict.json', 'r') as f:
                persona_dict = json.load(f)
            question_options = []
            new_question_regex = r"New Question:\s+(.*)\n"
            rationale_regex = r"Rationale:\s+(.*)"
            for i in range(10):
                new_question_output = question_developer(
                    original_question,
                    persona_dict["question_developer"],
                    temperature
                    )
                try:
                    question_only = re.search(new_question_regex, str(new_question_output)).group(1)
                    question_only = question_only.strip("\n")
                    question_options.append(new_question_output)
                    #print(f"Option {i+1}: {question_only}")
                except:
                    pass

            # Deduplicate the list of questions as the LLM sometimes repeats questions especially with a low temperature
            question_dedupe = []
            for question in question_options:
                if question not in question_dedupe:
                    question_dedupe.append(question) 
            for question in question_dedupe:
                question_only = re.search(new_question_regex, str(question)).group(1)
                output_placeholder.write(f"Option {question_dedupe.index(question)+1}: {question_only}")
            # Ask the user which question they prefer
            time.sleep(1)
            question_choice = int(input(f"Which question do you prefer? (1-{str(len(question_dedupe))})"))
            new_question = question_options[question_choice-1]
            output_placeholder.write(f"You chose option {question_choice}: {new_question}")

            new_question = re.search(new_question_regex, str(new_question_output)).group(1)
            rationale = re.search(rationale_regex, str(new_question_output)).group(1)

            summary.append(str((f"""Old Question: {original_question}
            New Question: {new_question}
            Rationale: {rationale}
            """)))
            summary_dict["Old Question"] = original_question
            summary_dict["New Question"] = new_question
            summary_dict["Rationale"] = rationale

            # Ask the LLM for a logical document title which is used as the pdf header later
            output_placeholder.write("Developing document title...")
            doctitle_question = question_developer(
                new_question+"\nDocument Title:",
                persona_dict["doctitle_developer"]
                )
            summary.append("Document Title: " + str(doctitle_question))
            summary_dict['Document Title'] = str(doctitle_question)

            # Ask the LLM for a logical file title which is used to name the files later
            output_placeholder.write("Developing file title...")    
            filetitle_question = question_developer(
                new_question+"\nFile Title:",
                persona_dict["filetitle_developer"]
                )
            summary.append("File Title: " + str(filetitle_question))
            summary_dict['File Title'] = str(filetitle_question)

            # get the best clinical question statement (PICO,PECO,SPIDER) from the llm
            output_placeholder.write("Determining which population statement template to use...")
            intorexp_question = question_developer(new_question, persona_dict["intorexp_developer"])
            summary.append(("Best clinical question statement: " + str(intorexp_question)))
            summary_dict["Best clinical question statement"] = str(intorexp_question)

            output_placeholder.write("Developing population statement...")
            #If interorexp contains intervention, then PICO is PICO
            # If the llm wants to use PICO or PECO, then create a PICO/PECO question
            if "pico" in intorexp_question.lower():
                pico_question = question_developer(new_question, persona_dict["PICO_developer"])
                summary.append("PICO: " + (pico_question))
                pop_statement = pico_question
            # If the llm wants to use SPIDER, then create a SPIDER question
            elif "peco" in intorexp_question.lower():
                pico_question = question_developer(new_question, persona_dict["PECO_developer"])
                summary.append("PECO Statement: "+(pico_question))
                pop_statement = pico_question
            # If the llm wants to use SPIDER, then create a SPIDER question
            elif "spider" in intorexp_question.lower():
                spider_question = question_developer(new_question, persona_dict["SPIDER_developer"])
                summary.append("SPIDER Statement: "+(spider_question))
                pop_statement = spider_question
            # Add Population Statement to summary dictionary
            summary_dict["Population Statement"] = pop_statement

            # Ask the LLM for inclusion and exclusion criteria
            output_placeholder.write("Developing inclusion and exclusion criteria...")
            incexc_question = question_developer(
                str(
                    pop_statement+
                    "\nMy Research Question: "+
                    new_question
                    ),
                persona_dict["incexc_developer"]
                )
            summary.append((incexc_question))
            summary_dict["Inclusion and Exclusion Criteria"] = incexc_question

            # Ask the LLM for a search strategy
            output_placeholder.write("Developing search strategy...")
            searchstrat_question = question_developer(
                str(
                    "My Research Question: "+
                    new_question+"\nPopulation Statement: "+
                    pop_statement+
                    "\n"+
                    incexc_question+
                    "\n\n"+
                    "Final Search Strategy:"
                ),
                persona_dict["searchstrat_developer"]
                )
            summary.append(("Final Search Strategy:\n"+searchstrat_question))
            summary_dict["Final Search Strategy"] = searchstrat_question

            # Ask the LLM for a list of databases to search
            output_placeholder.write("Determining which databases to search...")
            database_question = question_developer(
                str(new_question+"\nList of databases:"), persona_dict["database_developer"]
                )
            summary.append(("List of suggested databases:\n"+database_question))
            summary_dict["List of suggested databases"] = database_question
            output_placeholder.write("List of suggested databases:\n"+database_question)

            #convert whatever list the LLM returns into a list by asking the LLM to return a list
            database_list = question_developer(
                str(database_question+"\nList:('"), persona_dict["list_returner"]
                )
            if "and" in database_list:
                database_list = database_list.remove("and")
            database_list = eval(database_list)
            #print("Database list type = "+str(type(database_list)))
            #print("Database list = "+str(database_list))

            # Develop a search strategy for each database by asking the LLM for each database
            all_database_strats = ""
            output_placeholder.write(f"Developing search strategy for each database (n={len(database_list)})...")
            for database in database_list:
                output_placeholder.write("Developing search strategy for "+database+" ...")
                database_strat = question_developer(
                    str("My Search Strategy: "+
                        searchstrat_question+
                        "\nDatabase that will be searched: "+
                        database+"\nSearch strategy specific to "+
                        database+
                        ":"
                        ),
                    persona_dict["databasesearchstrat_developer"]
                    )
                summary.append(str("Search strategy specific to "+database+":\n"+database_strat))
                summary_dict["Search strategy specific to "+database] = database_strat
                print("Search strategy specific to "+database+":\n"+database_strat)
                all_database_strats += "\nStrategy for " + database + ":\n" + database_strat+"\n"

            summary_dict["All Database Strategy"] =(
                "Search Strategy for all databases:\n"+
                all_database_strats
            )
            
            output_placeholder.write("Writing file...")
            with open('summary_file.txt','w',encoding='utf-8') as handle:
                for item in summary:
                    handle.write(item+"\n")
            
            pubmed_question = question_developer(
                str(searchstrat_question)+"\n Pubmed Query: (", persona_dict["pubmedquery_developer"]
                )
            summary.append("Pubmed Query: "+str(pubmed_question))
            summary_dict["Pubmed Query"] = str(pubmed_question)

            docx_generator(summary_dict)
            return summary_dict
            questions = searchstrat_developer(user_input, output_placeholder=output_placeholder)

            st.markdown("## Choose your preferred research question")
            question_options = [re.search(r"New Question:\s+(.*)\n", q).group(1) for q in questions]

            selected_question = st.selectbox("Select a research question:", question_options)
            st.markdown(f"**Selected question:** {selected_question}")# Call your backend function here to generate the search strategy
        else:
            st.error("Please enter a research question.")
        pass

def question_developer(user_input, persona, temperature=0):
    """
    Module to generate a new research question based on the user input and persona
    """
    question = persona + user_input
    template = """Question: {question}

    Answer:"""
    prompt = PromptTemplate(template=template, input_variables=["question"])
    llm_chain = LLMChain(
        prompt=prompt,
        llm=ChatOpenAI(temperature=temperature, max_tokens=2000),
        verbose=False
        )


    return llm_chain.predict(question=question)

def searchstrat_developer(original_question, output_placeholder=None, temperature=0.3):
    """
    This is the main module for the generation of the search strategy
    The module takes the original question as input and returns a list 
    of the search strategy components which is then exported to pdf
    """
    summary = []
    summary_dict = {}

    # Ask the LLM for a new research question to-do: create a list of 5 different questions
    # Then ask the user which one they prefer or to keep the original question
    output_placeholder.write("Developing new research questions...")
    with open('perssonadict.json', 'r') as f:
        persona_dict = json.load(f)
    question_options = []
    new_question_regex = r"New Question:\s+(.*)\n"
    rationale_regex = r"Rationale:\s+(.*)"
    for i in range(10):
        new_question_output = question_developer(
            original_question,
            persona_dict["question_developer"],
            temperature
            )
        try:
            question_only = re.search(new_question_regex, str(new_question_output)).group(1)
            question_only = question_only.strip("\n")
            question_options.append(new_question_output)
            #print(f"Option {i+1}: {question_only}")
        except:
            pass

    # Deduplicate the list of questions as the LLM sometimes repeats questions especially with a low temperature
    question_dedupe = []
    for question in question_options:
        if question not in question_dedupe:
            question_dedupe.append(question) 
    for question in question_dedupe:
        question_only = re.search(new_question_regex, str(question)).group(1)
        output_placeholder.write(f"Option {question_dedupe.index(question)+1}: {question_only}")
    # Ask the user which question they prefer
    time.sleep(1)
    question_choice = int(input(f"Which question do you prefer? (1-{str(len(question_dedupe))})"))
    new_question = question_options[question_choice-1]
    output_placeholder.write(f"You chose option {question_choice}: {new_question}")

    new_question = re.search(new_question_regex, str(new_question_output)).group(1)
    rationale = re.search(rationale_regex, str(new_question_output)).group(1)

    summary.append(str((f"""Old Question: {original_question}
    New Question: {new_question}
    Rationale: {rationale}
    """)))
    summary_dict["Old Question"] = original_question
    summary_dict["New Question"] = new_question
    summary_dict["Rationale"] = rationale

    # Ask the LLM for a logical document title which is used as the pdf header later
    output_placeholder.write("Developing document title...")
    doctitle_question = question_developer(
        new_question+"\nDocument Title:",
        persona_dict["doctitle_developer"]
        )
    summary.append("Document Title: " + str(doctitle_question))
    summary_dict['Document Title'] = str(doctitle_question)

    # Ask the LLM for a logical file title which is used to name the files later
    output_placeholder.write("Developing file title...")    
    filetitle_question = question_developer(
        new_question+"\nFile Title:",
        persona_dict["filetitle_developer"]
        )
    summary.append("File Title: " + str(filetitle_question))
    summary_dict['File Title'] = str(filetitle_question)

    # get the best clinical question statement (PICO,PECO,SPIDER) from the llm
    output_placeholder.write("Determining which population statement template to use...")
    intorexp_question = question_developer(new_question, persona_dict["intorexp_developer"])
    summary.append(("Best clinical question statement: " + str(intorexp_question)))
    summary_dict["Best clinical question statement"] = str(intorexp_question)

    output_placeholder.write("Developing population statement...")
    #If interorexp contains intervention, then PICO is PICO
    # If the llm wants to use PICO or PECO, then create a PICO/PECO question
    if "pico" in intorexp_question.lower():
        pico_question = question_developer(new_question, persona_dict["PICO_developer"])
        summary.append("PICO: " + (pico_question))
        pop_statement = pico_question
    # If the llm wants to use SPIDER, then create a SPIDER question
    elif "peco" in intorexp_question.lower():
        pico_question = question_developer(new_question, persona_dict["PECO_developer"])
        summary.append("PECO Statement: "+(pico_question))
        pop_statement = pico_question
    # If the llm wants to use SPIDER, then create a SPIDER question
    elif "spider" in intorexp_question.lower():
        spider_question = question_developer(new_question, persona_dict["SPIDER_developer"])
        summary.append("SPIDER Statement: "+(spider_question))
        pop_statement = spider_question
    # Add Population Statement to summary dictionary
    summary_dict["Population Statement"] = pop_statement

    # Ask the LLM for inclusion and exclusion criteria
    output_placeholder.write("Developing inclusion and exclusion criteria...")
    incexc_question = question_developer(
        str(
            pop_statement+
            "\nMy Research Question: "+
            new_question
            ),
        persona_dict["incexc_developer"]
        )
    summary.append((incexc_question))
    summary_dict["Inclusion and Exclusion Criteria"] = incexc_question

    # Ask the LLM for a search strategy
    output_placeholder.write("Developing search strategy...")
    searchstrat_question = question_developer(
        str(
            "My Research Question: "+
            new_question+"\nPopulation Statement: "+
            pop_statement+
            "\n"+
            incexc_question+
            "\n\n"+
            "Final Search Strategy:"
        ),
        persona_dict["searchstrat_developer"]
        )
    summary.append(("Final Search Strategy:\n"+searchstrat_question))
    summary_dict["Final Search Strategy"] = searchstrat_question

    # Ask the LLM for a list of databases to search
    output_placeholder.write("Determining which databases to search...")
    database_question = question_developer(
        str(new_question+"\nList of databases:"), persona_dict["database_developer"]
        )
    summary.append(("List of suggested databases:\n"+database_question))
    summary_dict["List of suggested databases"] = database_question
    output_placeholder.write("List of suggested databases:\n"+database_question)

    #convert whatever list the LLM returns into a list by asking the LLM to return a list
    database_list = question_developer(
        str(database_question+"\nList:('"), persona_dict["list_returner"]
        )
    if "and" in database_list:
        database_list = database_list.remove("and")
    database_list = eval(database_list)
    #print("Database list type = "+str(type(database_list)))
    #print("Database list = "+str(database_list))

    # Develop a search strategy for each database by asking the LLM for each database
    all_database_strats = ""
    output_placeholder.write(f"Developing search strategy for each database (n={len(database_list)})...")
    for database in database_list:
        output_placeholder.write("Developing search strategy for "+database+" ...")
        database_strat = question_developer(
            str("My Search Strategy: "+
                searchstrat_question+
                "\nDatabase that will be searched: "+
                database+"\nSearch strategy specific to "+
                database+
                ":"
                ),
            persona_dict["databasesearchstrat_developer"]
            )
        summary.append(str("Search strategy specific to "+database+":\n"+database_strat))
        summary_dict["Search strategy specific to "+database] = database_strat
        print("Search strategy specific to "+database+":\n"+database_strat)
        all_database_strats += "\nStrategy for " + database + ":\n" + database_strat+"\n"

    summary_dict["All Database Strategy"] =(
        "Search Strategy for all databases:\n"+
        all_database_strats
    )
    
    output_placeholder.write("Writing file...")
    with open('summary_file.txt','w',encoding='utf-8') as handle:
        for item in summary:
            handle.write(item+"\n")
    
    pubmed_question = question_developer(
        str(searchstrat_question)+"\n Pubmed Query: (", persona_dict["pubmedquery_developer"]
        )
    summary.append("Pubmed Query: "+str(pubmed_question))
    summary_dict["Pubmed Query"] = str(pubmed_question)

    docx_generator(summary_dict)
    return summary_dict

def docx_generator(summary_dict):
    """
    The word_generator function takes the summary_dict dictionary and generates a Word 
    document with the summary information.
    It uses the templates/wordtemplate.docx for the formatting.
    """
    print("Generating Word summary of search strategy...")
    document_title = summary_dict["Document Title"]
    old_question = summary_dict["Old Question"]
    rationale = summary_dict["Rationale"]
    new_question = summary_dict["New Question"]
    best_clinical_question_statement = summary_dict["Best clinical question statement"]
    population_statement = summary_dict["Population Statement"]
    inclusion_exclusion_criteria = summary_dict["Inclusion and Exclusion Criteria"]
    final_search_strategy = summary_dict["Final Search Strategy"]
    suggested_databases = summary_dict["List of suggested databases"]
    all_database_strategy = summary_dict["All Database Strategy"]

    # Create a new Word document
    doc = docx.Document()

    # Set the document title
    doc.add_heading(str(document_title).strip('"'), 0)

    #add the logo
    doc.add_picture('synthscopelogo.png', width=Inches(4))

    # Add the summary information to the document
    doc.add_heading('Old Question', level=1)
    doc.add_paragraph(old_question)

    doc.add_heading('New Question', level=1)
    doc.add_paragraph(new_question)

    doc.add_heading('Rationale', level=1)
    doc.add_paragraph(rationale)

    doc.add_heading('Best Clinical Question Statement', level=1)
    doc.add_paragraph(best_clinical_question_statement)

    doc.add_heading('Population Statement', level=1)
    doc.add_paragraph(population_statement)

    doc.add_heading('Inclusion and Exclusion Criteria', level=1)
    doc.add_paragraph(inclusion_exclusion_criteria)

    doc.add_heading('Final Search Strategy', level=1)
    doc.add_paragraph(final_search_strategy)

    doc.add_heading('List of suggested databases', level=1)
    doc.add_paragraph(suggested_databases)

    doc.add_heading('All Database Strategy', level=1)
    doc.add_paragraph(all_database_strategy)

    file_name = str(summary_dict["File Title"]).strip('"').strip('.pdf') + ".docx"

    # Save the document
    doc.save('./search_summaries/' + file_name)
    time.sleep(1)
    convert('./search_summaries/' + file_name)
    print("File generated at search_summaries/"+file_name)


def database_search():
    st.header("Database Search")
    st.write("SynthScope will run the search strategy on a range of academic databases.")
    databases = st.multiselect("Select databases to search", ["Database 1", "Database 2", "Database 3"])
    if st.button("Run Database Search"):
        # Call your backend function here to run the database search
        pass


def title_abstract_screening():
    st.header("Title and Abstract Screening")
    st.write("SynthScope will screen the titles and abstracts of the results for if they meet the search strategy criteria and creates a file of articles that met criteria.")
    if st.button("Start Title and Abstract Screening"):
        # Call your backend function here to perform title and abstract screening
        pass
        uploaded_file = st.file_uploader("Choose a file to upload", type=['txt', 'csv', 'json'])

    if uploaded_file is not None:
        file_details = {
            "filename": uploaded_file.name,
            "filetype": uploaded_file.type,
            "file_size": uploaded_file.size
        }
        st.write(file_details)

        # Read the file content
        file_content = uploaded_file.getvalue().decode("utf-8")
        st.write(file_content)


def full_text_pdf_retrieval():
    st.header("Full Text PDF Retrieval")
    st.write("SynthScope will pull the full text pdf files for the papers that passed the first screening.")
    if st.button("Retrieve Full Text PDFs"):
        # Call your backend function here to retrieve full text PDFs
        pass


def full_text_screening():
    st.header("Full Text Screening")
    st.write("SynthScope will complete a screening of the full text of the papers and determines which ones meet the criteria.")
    if st.button("Start Full Text Screening"):
        # Call your backend function here to perform full text screening
        pass


def risk_of_bias_screening():
    st.header("Risk of Bias Screening")
    st.write("SynthScope will screen the included papers against the cochrane risk of bias tool and create a risk of bias table, risk of bias graph, and risk of biassummary.")
    if st.button("Start Risk of Bias Screening"):
    # Call your backend function here to perform risk of bias screening
        pass

def summary_generation():
    st.header("Summary Generation")
    st.write("SynthScope will summarize the full text papers that met criteria and return this summary of the articles that met criteria to the user.")
    if st.button("Generate Summary"):
    # Call your backend function here to generate the summary
        pass

def report_generation():
    st.header("Report Generation")
    st.write("SynthScope will create a full literature review report on the above process including introduction, methods, results, discussion, and conclusion.")
    report_format = st.selectbox("Choose report format", ["PDF", "Word", "HTML"])
    if st.button("Generate Report"):
    # Call your backend function here to generate the report
        pass

if __name__ == "__main__":
    main()
